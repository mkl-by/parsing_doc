from builtins import enumerate

from flask_login import login_required

import models, re, datetime
from flask import Blueprint, render_template, request
from pathlib import Path, PurePath
from path.forms import MyForm
from docx import Document
from app import db

paths=Blueprint('paths', __name__, template_folder='templates')
path_home = Path.home()
expansion=['.doc', '.docx']
dat=['номер документа', 'экземпляр', 'наименование документа', 'страницы']
pars=('дело', 'том')

def pars_data(text):
    #Написать проверку входящего и исходящего !!!!!!!!!
    in_reg=r'\d{,4}[c]{0,2}[d]{0,1}[c]{0,1}[p]{0,1}-cd' #входящий документ ^\d{,4}((c{,2}|dcp{1})-cd)$ \d{,4}((c{,2}|dcp{1})-cd) r'^\d{,4}-cd$'
    out_reg=r'^\d{,2}[/]\d[/]\d{,3}[c]{0,2}[d]{0,1}[c]{0,1}[p]{0,1}-cd$' #исходящий документ
    reg = r'^[0-3][0-9][.][0,1][0-2][.][1-2][0,9]\d\d$' # дата
    t=text.split('\n')
    p = re.compile(reg)
    p_in=re.compile(in_reg)

    p_out=re.compile(out_reg)

    try:
        #вернуть ошибку или подсветить ошибку в случае отсутствии даты!!!!!!!!!!!!!!!!!!!
        # print((p.findall(t[1])[0].split('.')))# сделать проверку на буквы!!!!!!!!!!!

        in_docum=(p_in.findall(t[0]))

        out_docum=(p_out.findall(t[0]))


        if in_docum==[] and out_docum==[]:
            in_out=None
        elif in_docum==[]:
            in_out=str(out_docum[0])
        else:
            in_out=str(in_docum[0])

        d=p.findall(t[1])[0].split('.')
        dat=datetime.date(int(d[2]), int(d[1]), int(d[0]))
    except IndexError:
        dat=None

    raspars={'data': dat, 'in_out_document':in_out}
    return raspars


def parsing_doc(path):
    q=['-','', ' ']
    d={}
    document = Document(path)
    for para in document.paragraphs:
        #выбираем все параграфы(строки) из документа и построчно читаем
        par=para.text
        num1=par.find('дело')
        num2 = par.find('том')
        if num1!=-1 or num2!=-1:
            #если в стоку входят слова "дело" или "том"
            spl=par.split()
            for st, spis in enumerate(spl):
                if spis.isdigit() and spl[st-1]=='дело':
                    case_number=int(spis)
                elif spis.isdigit() and spl[st-1]=='том':
                    case_tom=int(spis)

    for table in document.tables:
        #разбираем документ сначала на таблицы, потом на строки, потом поячеячно

        for s, row in enumerate(table.rows):

            sss=0
            do = models.Document()
            doc = list(do._sa_class_manager)

            for ss, i in enumerate(row.cells):
                if s >= 3 and not(i.text in q) and ss!=0:

                    do.__setattr__(doc[sss+1], i.text)
                    if ss==1 or ss==2:
                        dat=pars_data(i.text)
                        #возвращаем ошибку если в дате или номере документа ошибка
                        if dat['in_out_document']==None and dat['data']==None:
                            errors = {'number_row': s-2, 'error': 'ошибка в номере и дате документа, cтрока'}
                            return errors
                        elif dat['in_out_document']==None:
                            errors={'number_row': s-2, 'error':'ошибка в номере документа, строка'}
                            return errors
                        elif dat['data']==None:
                            errors = {'number_row': s-2, 'error': 'ошибка в дате документа, строка'}
                            return errors

                        do.__setattr__(doc[1], dat['in_out_document']) #номер документа
                        do.__setattr__(doc[8], dat['data'])
                    sss+=1
            do.__setattr__(doc[5], case_number)
            do.__setattr__(doc[6], case_tom)
            do.set_case()
            print(do.number, do.date_doc)
            # post = models.Document.query.filter(models.Document.number == do.number and models.Document.date_doc == do.date_doc)
            # print(post)
            #пишем информацию в базу
            if s >= 3: #т.к. первые 3 строки таблицы не данные
                post = models.Document.query.filter(models.Document.number == do.number and models.Document.date_doc == do.date_doc).count()
                print(post)
                if post==0:
            # для записи информации в базу
                    try:
                        db.session.add(do)
                        db.session.commit()
                    except:
                        db.session.rollback()
                        raise
                    finally:
                        db.session.close()

    return {'number_row':None, 'error':'Ошибок нет, документы добавлены в базу'}


@paths.route('/', methods=['GET', 'POST'])
@login_required
def my_path():
    erors=None
    form = MyForm()
    pren = str(path_home)
    if form.validate_on_submit():
        p=Path(request.form['form_s'])
        erors=parsing_doc(p)
        print(erors)


    directory = []
    files_dict={}

    path_isdir = Path.iterdir(path_home) # директория в которой ищем,
    p=request.args.get('urr')

    if p:
        path_isdir = Path.iterdir(Path(p))
        pren=str(Path(p).parent)# путь папка назад
    #[dd.append(x.name) for x in path_home.iterdir() if x.is_file() and (x.suffix in expansion)]

    for i in path_isdir:
        if i.is_file() and (i.suffix in expansion): #если файл и расширение ".doc", "docx","txt"
            files_dict[i.name]=i #добавляем в словарь имя файла как ключ и путь к нему
        elif i.is_dir() and not (i.parts[-1:][0].startswith('.')): #если директория и начинается на "."
            directory.append(i)

    return render_template('path/route.html', path_isdir=directory, form=form, pren=pren, files_dict=files_dict, errors=erors)
